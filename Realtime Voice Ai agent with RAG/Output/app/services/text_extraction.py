import os
from pypdf import PdfReader
from docx import Document as DocxDocument

class TextExtractionService:

    SUPPORTED_FORMATS = {
        "text/plain": ["txt", "md"],
        "application/pdf": ["pdf"],
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ["docx"],
    }

    def extract_text(self, file_path: str, content_type: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        ext = self._get_extension(file_path)
        if content_type == "text/plain" or ext in ["txt", "md"]:
            return self._extract_text_file(file_path)
        elif content_type == "application/pdf" or ext == "pdf":
            return self._extract_pdf(file_path)
        elif "wordprocessingml" in content_type or ext == "docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {content_type} ({ext})")

    def _extract_text_file(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read().strip()

    def _extract_pdf(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        parts = [p.extract_text() for p in reader.pages if p.extract_text()]
        return "\n\n".join(parts).strip()

    def _extract_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n\n".join(parts).strip()

    def is_supported(self, content_type: str, file_path: str) -> bool:
        ext = self._get_extension(file_path)
        if content_type in self.SUPPORTED_FORMATS:
            return True
        return any(ext in exts for exts in self.SUPPORTED_FORMATS.values())

    def _get_extension(self, file_path: str) -> str:
        return os.path.splitext(file_path)[1].lstrip(".").lower()
